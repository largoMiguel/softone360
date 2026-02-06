"""
Generador de Informes Institucionales para Plan de Desarrollo Municipal (PDM)
Basado en estándares colombianos de gestión pública territorial

Este módulo genera informes de gestión institucional alineados con:
- Constitución Política de Colombia de 1991
- Normatividad en planeación territorial
- Metodología General Ajustada (MGA) del DNP
- Agenda 2030 y Objetivos de Desarrollo Sostenible

Estructura del informe:
1. Portada institucional (con equipo de gobierno)
2. Introducción (marco legal y objetivo del informe)
3. Avance por líneas estratégicas (pilares del plan)
4. Avance por sectores MGA (áreas temáticas)
5. Avance por ODS (alineación con Agenda 2030)
6. Descripción de cumplimiento de metas
7. Ejecución del plan de acción por vigencia

Formato: PDF, DOCX, Excel
Estilo: Lenguaje técnico-administrativo, formal, tercera persona
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, Image as RLImage, KeepTogether
)
from io import BytesIO
from datetime import datetime
from typing import List, Dict, Any
import os
import base64
from collections import defaultdict

# Configurar matplotlib para uso en servidor (sin display)
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib import patches
import numpy as np
plt.rcParams['font.family'] = 'DejaVu Sans'

from sqlalchemy.orm import Session
from app.models.pdm import PdmActividadEvidencia
from app.models.pdm_ejecucion import PDMEjecucionPresupuestal
from app.models.user import User

class PDMReportGenerator:
    """Generador de informes PDF con estructura general con mejoras de rendimiento y contenido"""
    
    def __init__(self, entity, productos: List, actividades: List, anio: int, db: Session = None, filtros: dict = None, usar_ia: bool = False):
        self.entity = entity
        self.productos = productos
        self.actividades = actividades
        self.anio = anio
        self.db = db
        self.filtros = filtros or {}
        self.usar_ia = usar_ia  # Parámetro para habilitar resúmenes con IA
        self.buffer = BytesIO()
        self.doc = None
        self.styles = None
        self.story = []
        self.page_number = 0
        self._cache_graficas = {}  # Caché para evitar regenerar gráficas
        
    def get_justify_style(self, fontSize=8):
        """Helper para crear estilos justificados reutilizables"""
        return ParagraphStyle(
            'JustifyStyle',
            parent=self.styles['Normal'],
            alignment=TA_JUSTIFY,
            fontSize=fontSize,
            leading=fontSize + 2
        )
        
    def add_header_footer(self, canvas, doc):
        """Encabezado y pie de página estándar"""
        canvas.saveState()
        
        # ENCABEZADO
        canvas.setFont('Helvetica', 8)
        # Código de formulario estándar
        canvas.drawString(0.5*inch, 10.5*inch, "FM-PDM-001")
        canvas.drawString(0.5*inch, 10.3*inch, "Versión: 1.0")
        
        # Número de página y título
        canvas.drawRightString(8*inch, 10.5*inch, f"Página {doc.page}")
        canvas.drawRightString(8*inch, 10.3*inch, "INFORME DE GESTIÓN INSTITUCIONAL")
        
        # Línea separadora
        canvas.setStrokeColor(colors.HexColor('#003366'))
        canvas.line(0.5*inch, 10.2*inch, 8*inch, 10.2*inch)
        
        # PIE DE PÁGINA
        canvas.setFont('Helvetica', 7)
        footer_text = f"Plan de Desarrollo Municipal - {self.entity.name}"
        canvas.drawCentredString(4.25*inch, 0.5*inch, footer_text)
        
        canvas.restoreState()
    
    def generate_portada(self):
        """Genera la portada institucional según formato colombiano"""
        # Título principal institucional
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#003366'),
            alignment=TA_CENTER,
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Spacer(1, 1.5*inch))
        self.story.append(Paragraph("INFORME DE GESTIÓN INSTITUCIONAL", title_style))
        self.story.append(Spacer(1, 0.2*inch))
        
        # Año de vigencia
        anio_texto = "Vigencia 2024-2027" if self.anio == 0 else f"Vigencia {self.anio}"
        anio_style = ParagraphStyle(
            'AnioStyle',
            parent=self.styles['Heading2'],
            fontSize=18,
            textColor=colors.HexColor('#003366'),
            alignment=TA_CENTER,
            spaceAfter=20
        )
        self.story.append(Paragraph(anio_texto, anio_style))
        self.story.append(Spacer(1, 0.3*inch))
        
        # Nombre del plan
        plan_style = ParagraphStyle(
            'PlanTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#003366'),
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        plan_name = getattr(self.entity, 'plan_name', None) or "PLAN DE DESARROLLO MUNICIPAL"
        self.story.append(Paragraph(plan_name, plan_style))
        self.story.append(Spacer(1, 0.2*inch))
        
        # Tipo de informe
        tipo_style = ParagraphStyle(
            'TipoInforme',
            parent=self.styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#666666'),
            spaceAfter=6
        )
        self.story.append(Paragraph("Informe de Gestión / Rendición de Cuentas", tipo_style))
        self.story.append(Spacer(1, 0.4*inch))
        
        # Entidad - Municipio
        entity_style = ParagraphStyle(
            'EntityName',
            parent=self.styles['Heading2'],
            fontSize=16,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#003366'),
            fontName='Helvetica-Bold'
        )
        self.story.append(Paragraph(self.entity.name.upper(), entity_style))
        
        alcaldia_style = ParagraphStyle(
            'Alcaldia',
            parent=self.styles['Normal'],
            fontSize=14,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#666666'),
            spaceAfter=12
        )
        self.story.append(Paragraph("Alcaldía Municipal", alcaldia_style))
        self.story.append(Spacer(1, 0.4*inch))
        
        # Información de filtros si existen
        if self.filtros:
            filter_info = []
            if self.filtros.get('secretarias'):
                secs = ', '.join(self.filtros['secretarias'])
                filter_info.append(f"Secretarías: {secs}")
            if self.filtros.get('fecha_inicio') or self.filtros.get('fecha_fin'):
                inicio = self.filtros.get('fecha_inicio', 'N/A')
                fin = self.filtros.get('fecha_fin', 'N/A')
                filter_info.append(f"Período: {inicio} a {fin}")
            if self.filtros.get('estados'):
                estados = ', '.join(self.filtros['estados'])
                filter_info.append(f"Estados: {estados}")
            
            if filter_info:
                filter_style = ParagraphStyle(
                    'FilterInfo',
                    parent=self.styles['Normal'],
                    fontSize=10,
                    alignment=TA_CENTER,
                    textColor=colors.HexColor('#666666'),
                    spaceAfter=6
                )
                self.story.append(Spacer(1, 0.3*inch))
                for info in filter_info:
                    self.story.append(Paragraph(info, filter_style))
        
        # Equipo de gobierno (si está disponible)
        equipo_style = ParagraphStyle(
            'EquipoGobierno',
            parent=self.styles['Normal'],
            fontSize=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor('#666666'),
            spaceAfter=4,
            leading=14
        )
        
        self.story.append(Spacer(1, 0.3*inch))
        self.story.append(Paragraph("<b>Equipo de Gobierno Municipal</b>", equipo_style))
        
        # Información del equipo (placeholder - puede venir de la entidad)
        equipo_info = [
            "Alcalde Municipal",
            "Gestor(a) Social",
            "Jefe de Planeación Municipal",
            "Secretario de Gobierno",
            "Comisaría de Familia",
            "Inspector de Policía"
        ]
        
        for cargo in equipo_info:
            self.story.append(Paragraph(cargo, equipo_style))
        
        self.story.append(PageBreak())
    
    def generate_introduccion(self):
        """Genera la introducción institucional del informe"""
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#003366'),
            spaceAfter=16,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Paragraph("INTRODUCCIÓN", title_style))
        
        anio_texto = "el cuatrienio 2024-2027" if self.anio == 0 else f"la vigencia {self.anio}"
        plan_name = getattr(self.entity, 'plan_name', 'Plan de Desarrollo Municipal')
        
        intro_text = f"""
        Los planes de desarrollo de las entidades territoriales constituyen la carta de navegación y el principal 
        instrumento de planeación para el desarrollo integral del territorio. Se configuran como herramientas de 
        carácter político y técnico, construidas mediante procesos democráticos y pluralistas, en las cuales se 
        materializan las decisiones, acciones, medios y recursos que orientan la gestión pública hacia el logro de 
        los objetivos de desarrollo territorial.
        <br/><br/>
        El Plan de Desarrollo Municipal "{plan_name}" fue adoptado mediante Acuerdo Municipal, en cumplimiento de 
        lo establecido en la Constitución Política de Colombia de 1991 y la normatividad vigente en materia de 
        planeación territorial. Este instrumento define las estrategias, programas y proyectos que guían la acción 
        gubernamental del municipio de {self.entity.name}.
        <br/><br/>
        El presente informe de gestión institucional tiene como objetivo presentar un balance integral de los 
        resultados alcanzados durante {anio_texto}, evidenciando el estado de ejecución de las metas programadas, 
        la inversión de recursos administrativos y financieros, así como el avance en el cumplimiento de los 
        compromisos adquiridos con la comunidad.
        <br/><br/>
        Este documento describe los logros y avances obtenidos, identifica las metas pendientes de cumplimiento, 
        y formula recomendaciones estratégicas para el fortalecimiento de la gestión pública municipal. La información 
        contenida se encuentra organizada por líneas estratégicas, sectores de intervención y su alineación con los 
        Objetivos de Desarrollo Sostenible (ODS) de la Agenda 2030.
        <br/><br/>
        El informe se estructura como un instrumento de rendición de cuentas ante la comunidad y de transparencia en 
        la gestión integral del territorio, enmarcado en los principios de eficiencia, eficacia y efectividad de la 
        administración pública.
        """
        
        justify_style = ParagraphStyle(
            'Justify',
            parent=self.styles['BodyText'],
            alignment=TA_JUSTIFY,
            fontSize=10
        )
        
        self.story.append(Paragraph(intro_text, justify_style))
        
        # RESUMEN EJECUTIVO con KPIs principales
        self.generar_resumen_ejecutivo()
        
        # RESUMEN CON IA (si está habilitado)
        if self.usar_ia:
            self.generar_resumen_ia()
        
        self.story.append(PageBreak())
    
    def generar_resumen_ejecutivo(self):
        """Genera resumen ejecutivo con indicadores clave al inicio del informe"""
        try:
            title_style = ParagraphStyle(
                'ExecutiveTitle',
                parent=self.styles['Heading1'],
                fontSize=14,
                textColor=colors.HexColor('#003366'),
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            
            self.story.append(Spacer(1, 0.2*inch))
            self.story.append(Paragraph("RESUMEN EJECUTIVO", title_style))
            
            # Calcular KPIs generales
            total_productos = len(self.productos)
            # Total de actividades según filtro de año
            if self.anio == 0:
                total_actividades = len(self.actividades)
            else:
                total_actividades = sum(1 for act in self.actividades if act.anio == self.anio)
            
            # Avance promedio (solo considerar productos con programación en el año actual)
            suma_avances = 0
            productos_con_meta = 0
            total_meta = 0
            total_ejecutado = 0
            
            for prod in self.productos:
                # Obtener meta programada según año
                if self.anio == 0:
                    # Sumar todas las metas del cuatrienio
                    meta_anio = (
                        (getattr(prod, 'programacion_2024', 0) or 0) +
                        (getattr(prod, 'programacion_2025', 0) or 0) +
                        (getattr(prod, 'programacion_2026', 0) or 0) +
                        (getattr(prod, 'programacion_2027', 0) or 0)
                    )
                else:
                    # Solo meta del año específico
                    meta_anio = getattr(prod, f'programacion_{self.anio}', 0) or 0
                
                if meta_anio > 0:
                    productos_con_meta += 1
                    total_meta += meta_anio
                    suma_avances += self.calcular_avance_producto(prod)
                
            avance_promedio = suma_avances / productos_con_meta if productos_con_meta > 0 else 0
            
            # Avance financiero promedio (sobre todos los productos)
            suma_financiero = 0
            for prod in self.productos:
                suma_financiero += self.calcular_avance_financiero(prod)
            avance_financiero_promedio = suma_financiero / total_productos if total_productos > 0 else 0
            
            # Actividades por estado (según año filtrado o todas si anio=0)
            estados_count = {}
            for act in self.actividades:
                # Si anio es 0, incluir todas las actividades
                if self.anio == 0 or act.anio == self.anio:
                    estado = act.estado
                    estados_count[estado] = estados_count.get(estado, 0) + 1
            
            # Total presupuesto según año seleccionado
            total_presupuesto = 0
            for prod in self.productos:
                if self.anio == 0:
                    # Presupuesto del cuatrienio completo (suma de todos los años)
                    total_presupuesto += float(prod.total_2024 or 0)
                    total_presupuesto += float(prod.total_2025 or 0)
                    total_presupuesto += float(prod.total_2026 or 0)
                    total_presupuesto += float(prod.total_2027 or 0)
                else:
                    # Presupuesto del año específico
                    if self.anio == 2024:
                        total_presupuesto += float(prod.total_2024 or 0)
                    elif self.anio == 2025:
                        total_presupuesto += float(prod.total_2025 or 0)
                    elif self.anio == 2026:
                        total_presupuesto += float(prod.total_2026 or 0)
                    elif self.anio == 2027:
                        total_presupuesto += float(prod.total_2027 or 0)
            
            # TABLA DE KPIs PRINCIPALES
            white_bold = ParagraphStyle('WhiteBold', parent=self.styles['Normal'], 
                                       textColor=colors.white, fontName='Helvetica-Bold', fontSize=9)
            center_style = ParagraphStyle('Center', parent=self.styles['Normal'], 
                                         alignment=TA_CENTER, fontSize=10)
            
            kpis_data = [
                [
                    Paragraph('Total Productos', white_bold),
                    Paragraph('Avance Físico Promedio', white_bold),
                    Paragraph('Avance Financiero Promedio', white_bold),
                    Paragraph('Presupuesto Acumulado', white_bold)
                ],
                [
                    Paragraph(f'<b>{total_productos}</b>', center_style),
                    Paragraph(f'<b>{avance_promedio:.1f}%</b>', center_style),
                    Paragraph(f'<b>{avance_financiero_promedio:.1f}%</b>', center_style),
                    Paragraph(f'<b>${total_presupuesto:,.0f}</b>', center_style)
                ]
            ]
            
            kpis_table = Table(kpis_data, colWidths=[1.75*inch, 1.75*inch, 1.75*inch, 1.75*inch])
            kpis_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#E8F4F8')),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#003366')),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            self.story.append(kpis_table)
            self.story.append(Spacer(1, 0.15*inch))
            
            # TABLA DE ACTIVIDADES POR ESTADO
            if estados_count:
                actividades_data = [
                    [Paragraph('Estado de Actividades', white_bold), 
                     Paragraph('Cantidad', white_bold),
                     Paragraph('Porcentaje', white_bold)]
                ]
                
                total_act_anio = sum(estados_count.values())
                for estado, count in sorted(estados_count.items()):
                    porcentaje = (count / total_act_anio * 100) if total_act_anio > 0 else 0
                    actividades_data.append([
                        Paragraph(estado, self.styles['Normal']),
                        Paragraph(f'{count}', center_style),
                        Paragraph(f'{porcentaje:.1f}%', center_style)
                    ])
                
                act_table = Table(actividades_data, colWidths=[3*inch, 2*inch, 2*inch])
                act_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F5F5')),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                
                self.story.append(act_table)
            
            print("✅ Resumen ejecutivo generado")
            
        except Exception as e:
            print(f"⚠️ Error generando resumen ejecutivo: {e}")
            import traceback
            traceback.print_exc()
    
    def generar_resumen_ia(self):
        """
        Genera un resumen narrativo con IA (OpenAI) sobre el estado del PDM
        Mejora implementada: análisis inteligente opcional
        """
        try:
            import os
            from openai import OpenAI
            
            title_style = ParagraphStyle(
                'AITitle',
                parent=self.styles['Heading1'],
                fontSize=14,
                textColor=colors.HexColor('#003366'),
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            
            self.story.append(Spacer(1, 0.2*inch))
            self.story.append(Paragraph("ANÁLISIS NARRATIVO CON INTELIGENCIA ARTIFICIAL", title_style))
            
            # Preparar datos para el prompt
            total_productos = len(self.productos)
            total_actividades = len([a for a in self.actividades if self.anio == 0 or a.anio == self.anio])
            actividades_completadas = len([a for a in self.actividades if (self.anio == 0 or a.anio == self.anio) and a.estado == 'COMPLETADA'])
            
            suma_avances = sum(self.calcular_avance_producto(p) for p in self.productos)
            avance_promedio = suma_avances / total_productos if total_productos > 0 else 0
            
            # Crear prompt para OpenAI
            anio_texto = "todos los años del cuatrienio 2024-2027" if self.anio == 0 else f"el año {self.anio}"
            
            prompt = f"""Eres un analista experto en gestión pública territorial colombiana. 
            
Genera un análisis narrativo profesional y técnico del siguiente Plan de Desarrollo Municipal:

DATOS DEL INFORME:
- Entidad: {self.entity.name}
- Período: {anio_texto}
- Total de productos: {total_productos}
- Total de actividades: {total_actividades}
- Actividades completadas: {actividades_completadas} ({actividades_completadas/total_actividades*100 if total_actividades > 0 else 0:.1f}%)
- Avance físico promedio: {avance_promedio:.1f}%

El análisis debe:
1. Evaluar el nivel de cumplimiento general (excelente, bueno, regular, bajo)
2. Identificar fortalezas principales
3. Señalar áreas de mejora o riesgos
4. Dar recomendaciones estratégicas

Límite: 250 palabras. Usa lenguaje formal y técnico apropiado para gestión pública."""

            # Llamar a OpenAI
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Eres un experto en análisis de gestión pública territorial en Colombia."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.7
            )
            
            resumen_ia = response.choices[0].message.content
            
            # Agregar resumen al informe
            ia_style = ParagraphStyle(
                'IAText',
                parent=self.styles['BodyText'],
                alignment=TA_JUSTIFY,
                fontSize=10,
                spaceAfter=12,
                leftIndent=12,
                rightIndent=12,
                backColor=colors.HexColor('#F0F8FF'),
                borderPadding=10
            )
            
            self.story.append(Paragraph(f"<i>{resumen_ia}</i>", ia_style))
            self.story.append(Spacer(1, 0.15*inch))
            
            # Nota al pie
            nota_style = ParagraphStyle(
                'NotaIA',
                parent=self.styles['Normal'],
                fontSize=8,
                textColor=colors.HexColor('#666666'),
                alignment=TA_CENTER
            )
            self.story.append(Paragraph(
                "<i>* Análisis generado con Inteligencia Artificial (OpenAI GPT-4). "
                "Este resumen es orientativo y debe ser validado por el equipo técnico.</i>",
                nota_style
            ))
            
            print("✅ Resumen con IA generado exitosamente")
            
        except ImportError:
            print("⚠️ OpenAI no está instalado. Saltando resumen con IA.")
        except Exception as e:
            print(f"⚠️ Error generando resumen con IA: {e}")
            # No lanzar excepción, solo advertir
            import traceback
            traceback.print_exc()
    
    def calcular_avance_producto(self, producto):
        """
        Calcula el avance de un producto basado en meta ejecutada vs meta programada
        Respeta el año seleccionado (self.anio). Si anio=0, calcula promedio de todos los años.
        """
        try:
            # Determinar años a calcular según self.anio
            if self.anio == 0:
                # Todos los años del cuatrienio
                anios = [2024, 2025, 2026, 2027]
            else:
                # Solo el año seleccionado
                anios = [self.anio]
            
            suma_avances = 0
            total_anios_con_meta = 0
            
            print(f"\n🔍 Calculando avance para producto: {producto.codigo_producto}")
            print(f"   Año(s) a calcular: {anios}")
            print(f"   Total actividades disponibles: {len(self.actividades)}")
            
            for anio in anios:
                # Obtener meta programada del año
                meta_programada = getattr(producto, f'programacion_{anio}', 0) or 0
                
                if meta_programada > 0:
                    # Calcular meta ejecutada: suma de meta_ejecutar de actividades con evidencia
                    actividades_anio = [
                        act for act in self.actividades 
                        if act.codigo_producto == producto.codigo_producto and act.anio == anio
                    ]
                    
                    print(f"   Año {anio}: meta_programada={meta_programada}, actividades={len(actividades_anio)}")
                    
                    # Sumar meta_ejecutar de actividades que tienen evidencia
                    # Usar el flag tiene_evidencia agregado en el router (optimización para no cargar objetos)
                    meta_ejecutada = sum(
                        act.meta_ejecutar for act in actividades_anio 
                        if hasattr(act, 'tiene_evidencia') and act.tiene_evidencia
                    )
                    
                    actividades_con_evidencia = sum(
                        1 for act in actividades_anio 
                        if hasattr(act, 'tiene_evidencia') and act.tiene_evidencia
                    )
                    print(f"   Año {anio}: meta_ejecutada={meta_ejecutada}, actividades_con_evidencia={actividades_con_evidencia}")
                    
                    # Calcular porcentaje de avance (topar en 100%)
                    porcentaje_avance = min(100, (meta_ejecutada / meta_programada) * 100)
                    print(f"   Año {anio}: porcentaje_avance={porcentaje_avance:.1f}%")
                    suma_avances += porcentaje_avance
                    total_anios_con_meta += 1
            
            resultado = suma_avances / total_anios_con_meta if total_anios_con_meta > 0 else 0
            print(f"   ✅ Avance promedio final: {resultado:.1f}%\n")
            return resultado
            
        except Exception as e:
            print(f"      ⚠️ Error calculando avance para {producto.codigo_producto}: {e}")
            import traceback
            traceback.print_exc()
            return 0
    
    def calcular_avance_financiero(self, producto) -> float:
        """
        Calcula el avance financiero real basado en la ejecución presupuestal
        Formula: (Pagos / Presupuesto Definitivo) * 100
        
        Si no hay datos de ejecución, retorna el avance físico como estimación
        """
        try:
            if not self.db:
                # Sin acceso a DB, usar avance físico
                return self.calcular_avance_producto(producto)
            
            # Consultar ejecución presupuestal para este producto y año
            if self.anio == 0:
                # Todos los años del cuatrienio
                ejecuciones = self.db.query(PDMEjecucionPresupuestal).filter(
                    PDMEjecucionPresupuestal.entity_id == self.entity.id,
                    PDMEjecucionPresupuestal.codigo_producto == producto.codigo_producto
                ).all()
            else:
                # Solo año específico
                ejecuciones = self.db.query(PDMEjecucionPresupuestal).filter(
                    PDMEjecucionPresupuestal.entity_id == self.entity.id,
                    PDMEjecucionPresupuestal.codigo_producto == producto.codigo_producto,
                    PDMEjecucionPresupuestal.anio == self.anio
                ).all()
            
            if not ejecuciones:
                anio_texto = "todos los años" if self.anio == 0 else str(self.anio)
                print(f"      ℹ️ No hay ejecución presupuestal para {producto.codigo_producto} en {anio_texto}")
                # Sin datos de ejecución, usar avance físico como estimación
                return self.calcular_avance_producto(producto)
            
            # Sumar totales de todas las fuentes
            total_definitivo = 0
            total_pagos = 0
            
            for ejecucion in ejecuciones:
                total_definitivo += float(ejecucion.pto_definitivo or 0)
                total_pagos += float(ejecucion.pagos or 0)
            
            # Calcular porcentaje
            if total_definitivo == 0:
                print(f"      ⚠️ Presupuesto definitivo = 0 para {producto.codigo_producto}")
                return self.calcular_avance_producto(producto)
            
            avance_financiero = (total_pagos / total_definitivo) * 100
            print(f"      💰 Avance financiero {producto.codigo_producto}: {avance_financiero:.1f}% (Pagos: ${total_pagos:,.0f} / Definitivo: ${total_definitivo:,.0f})")
            
            return min(100, max(0, avance_financiero))  # Entre 0 y 100%
            
        except Exception as e:
            print(f"      ❌ Error calculando avance financiero para {producto.codigo_producto}: {e}")
            import traceback
            traceback.print_exc()
            # En caso de error, usar avance físico
            return self.calcular_avance_producto(producto)
    
    def generate_grafica_moderna_lineas(self):
        """Genera gráfica moderna de avance por líneas estratégicas con caché"""
        # Verificar caché (incluir año en la key)
        cache_key = f'grafica_lineas_{self.anio}'
        if cache_key in self._cache_graficas:
            print(f"   ⚡ Usando gráfica en caché: {cache_key}")
            self.story.append(self._cache_graficas[cache_key])
            self.story.append(Spacer(1, 0.3*inch))
            return
        
        # Calcular avance por línea (ya respeta self.anio por calcular_avance_producto)
        lineas_data = {}
        for prod in self.productos:
            linea = prod.linea_estrategica or 'Sin Línea'
            if linea not in lineas_data:
                lineas_data[linea] = {'total': 0, 'suma_avance': 0}
            lineas_data[linea]['total'] += 1
            avance = self.calcular_avance_producto(prod)
            lineas_data[linea]['suma_avance'] += avance
        
        lineas = []
        avances = []
        for linea, data in lineas_data.items():
            if data['total'] > 0:
                promedio = data['suma_avance'] / data['total']
                lineas.append(linea[:40])
                avances.append(promedio)
        
        if not lineas:
            return
        
        try:
            # Diseño moderno con colores institucionales
            fig, ax = plt.subplots(figsize=(9, max(len(lineas) * 0.6, 4)))
            fig.patch.set_facecolor('white')
            
            # Colores: verde institucional y gradientes
            colors = ['#4F9A54' if a >= 70 else '#FFA726' if a >= 50 else '#EF5350' for a in avances]
            
            y_pos = np.arange(len(lineas))
            bars = ax.barh(y_pos, avances, color=colors, height=0.6, alpha=0.9)
            
            # Agregar valores al final de cada barra
            for i, (bar, val) in enumerate(zip(bars, avances)):
                width = bar.get_width()
                ax.text(width + 2, bar.get_y() + bar.get_height()/2, 
                       f'{val:.1f}%', ha='left', va='center', 
                       fontsize=10, fontweight='bold', color='#333')
            
            ax.set_yticks(y_pos)
            ax.set_yticklabels(lineas, fontsize=9)
            ax.set_xlabel('Porcentaje de Avance (%)', fontsize=11, fontweight='bold', color='#333')
            ax.set_title('Avance por Línea Estratégica', fontsize=13, fontweight='bold', 
                        color='#003366', pad=20)
            ax.set_xlim(0, 110)
            
            # Estilo moderno
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#CCCCCC')
            ax.spines['bottom'].set_color('#CCCCCC')
            ax.grid(axis='x', alpha=0.2, linestyle='--', color='#CCCCCC')
            ax.set_axisbelow(True)
            
            plt.tight_layout()
            
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            img_buffer.seek(0)
            plt.close(fig)
            
            img = RLImage(img_buffer, width=7*inch, height=max(len(lineas) * 0.6*inch, 3.5*inch))
            
            # Guardar en caché
            self._cache_graficas[cache_key] = img
            print(f"   💾 Gráfica guardada en caché: {cache_key}")
            
            self.story.append(img)
            self.story.append(Spacer(1, 0.3*inch))
            
        except Exception as e:
            print(f"   ❌ Error generando gráfica de líneas: {str(e)}")
        finally:
            plt.close('all')
    
    def generate_grafica_moderna_sectores(self):
        """Genera gráfica moderna de avance por sectores MGA"""
        sectores_data = defaultdict(lambda: {'total': 0, 'suma_avance': 0})
        
        for prod in self.productos:
            sector = prod.sector_mga or 'Sin Sector'
            sectores_data[sector]['total'] += 1
            avance = self.calcular_avance_producto(prod)
            sectores_data[sector]['suma_avance'] += avance
        
        sectores = []
        avances = []
        for sector, data in sectores_data.items():
            if data['total'] > 0:
                promedio = data['suma_avance'] / data['total']
                sectores.append(sector[:40])
                avances.append(promedio)
        
        if not sectores:
            return
        
        try:
            fig, ax = plt.subplots(figsize=(9, max(len(sectores) * 0.6, 4)))
            fig.patch.set_facecolor('white')
            
            colors = ['#4F9A54' if a >= 70 else '#FFA726' if a >= 50 else '#EF5350' for a in avances]
            
            y_pos = np.arange(len(sectores))
            bars = ax.barh(y_pos, avances, color=colors, height=0.6, alpha=0.9)
            
            for i, (bar, val) in enumerate(zip(bars, avances)):
                width = bar.get_width()
                ax.text(width + 2, bar.get_y() + bar.get_height()/2, 
                       f'{val:.1f}%', ha='left', va='center', 
                       fontsize=10, fontweight='bold', color='#333')
            
            ax.set_yticks(y_pos)
            ax.set_yticklabels(sectores, fontsize=9)
            ax.set_xlabel('Porcentaje de Avance (%)', fontsize=11, fontweight='bold', color='#333')
            ax.set_title('Avance por Sector MGA', fontsize=13, fontweight='bold', 
                        color='#003366', pad=20)
            ax.set_xlim(0, 110)
            
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#CCCCCC')
            ax.spines['bottom'].set_color('#CCCCCC')
            ax.grid(axis='x', alpha=0.2, linestyle='--', color='#CCCCCC')
            ax.set_axisbelow(True)
            
            plt.tight_layout()
            
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            img_buffer.seek(0)
            plt.close(fig)
            
            img = RLImage(img_buffer, width=7*inch, height=max(len(sectores) * 0.6*inch, 3.5*inch))
            
            # Guardar en caché
            self._cache_graficas[cache_key] = img
            print(f"   💾 Gráfica guardada en caché: {cache_key}")
            
            self.story.append(img)
            self.story.append(Spacer(1, 0.3*inch))
            
        except Exception as e:
            print(f"   ❌ Error generando gráfica de sectores: {str(e)}")
        finally:
            plt.close('all')
    
    def generate_grafica_moderna_ods(self):
        """Genera gráfica moderna de avance por ODS con caché"""
        # Verificar caché (incluir año en la key)
        cache_key = f'grafica_ods_{self.anio}'
        if cache_key in self._cache_graficas:
            print(f"   ⚡ Usando gráfica en caché: {cache_key}")
            self.story.append(self._cache_graficas[cache_key])
            self.story.append(Spacer(1, 0.3*inch))
            return
        
        ods_data = defaultdict(lambda: {'total': 0, 'suma_avance': 0})
        
        for prod in self.productos:
            ods = prod.ods or 'Sin ODS'
            ods_data[ods]['total'] += 1
            avance = self.calcular_avance_producto(prod)
            ods_data[ods]['suma_avance'] += avance
        
        ods_list = []
        avances = []
        for ods, data in ods_data.items():
            if data['total'] > 0:
                promedio = data['suma_avance'] / data['total']
                ods_list.append(ods[:45])
                avances.append(promedio)
        
        if not ods_list:
            return
        
        try:
            fig, ax = plt.subplots(figsize=(9, max(len(ods_list) * 0.6, 4)))
            fig.patch.set_facecolor('white')
            
            colors = ['#4F9A54' if a >= 70 else '#FFA726' if a >= 50 else '#EF5350' for a in avances]
            
            y_pos = np.arange(len(ods_list))
            bars = ax.barh(y_pos, avances, color=colors, height=0.6, alpha=0.9)
            
            for i, (bar, val) in enumerate(zip(bars, avances)):
                width = bar.get_width()
                ax.text(width + 2, bar.get_y() + bar.get_height()/2, 
                       f'{val:.1f}%', ha='left', va='center', 
                       fontsize=10, fontweight='bold', color='#333')
            
            ax.set_yticks(y_pos)
            ax.set_yticklabels(ods_list, fontsize=9)
            ax.set_xlabel('Porcentaje de Avance (%)', fontsize=11, fontweight='bold', color='#333')
            ax.set_title('Avance por Objetivos de Desarrollo Sostenible', fontsize=13, 
                        fontweight='bold', color='#003366', pad=20)
            ax.set_xlim(0, 110)
            
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#CCCCCC')
            ax.spines['bottom'].set_color('#CCCCCC')
            ax.grid(axis='x', alpha=0.2, linestyle='--', color='#CCCCCC')
            ax.set_axisbelow(True)
            
            plt.tight_layout()
            
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
            img_buffer.seek(0)
            plt.close(fig)
            
            img = RLImage(img_buffer, width=7*inch, height=max(len(ods_list) * 0.6*inch, 3.5*inch))
            
            # Guardar en caché
            self._cache_graficas[cache_key] = img
            print(f"   💾 Gráfica guardada en caché: {cache_key}")
            
            self.story.append(img)
            self.story.append(Spacer(1, 0.3*inch))
            
        except Exception as e:
            print(f"   ❌ Error generando gráfica de ODS: {str(e)}")
        finally:
            plt.close('all')
    
    def generate_seccion_lineas(self):
        """Genera sección de avance por líneas estratégicas"""
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Paragraph(
            "AVANCE DE CUMPLIMIENTO DE METAS DEL PLAN DE DESARROLLO POR LÍNEAS ESTRATÉGICAS",
            title_style
        ))
        
        # Definición conceptual
        justify_style = ParagraphStyle(
            'Justify',
            parent=self.styles['BodyText'],
            alignment=TA_JUSTIFY,
            fontSize=10,
            spaceAfter=12
        )
        
        concepto_lineas = """
        Las líneas estratégicas constituyen los pilares, ejes o dimensiones fundamentales sobre los cuales 
        se estructura el Plan de Desarrollo Municipal. Estas líneas orientan la gestión pública y la asignación 
        de recursos hacia el logro de resultados específicos en áreas prioritarias del desarrollo territorial, 
        garantizando coherencia entre los objetivos de gobierno y las necesidades de la población.
        """
        
        self.story.append(Paragraph(concepto_lineas, justify_style))
        self.story.append(Spacer(1, 0.2*inch))
        
        # Generar gráfica moderna
        self.generate_grafica_moderna_lineas()
        
        # Generar tablas de productos por línea estratégica
        self.generate_tabla_productos()
        
        self.story.append(PageBreak())
    
    def generate_tabla_productos(self):
        """Genera tabla detallada de productos por línea estratégica"""
        # Agrupar productos por línea
        productos_por_linea = {}
        for prod in self.productos:
            linea = prod.linea_estrategica or 'Sin Línea Estratégica'
            if linea not in productos_por_linea:
                productos_por_linea[linea] = []
            productos_por_linea[linea].append(prod)
        
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Paragraph(
            "DESCRIPCIÓN DE CUMPLIMIENTO DE METAS PLAN DE DESARROLLO POR LÍNEAS ESTRATÉGICAS",
            title_style
        ))
        self.story.append(Spacer(1, 0.2*inch))
        
        for linea, productos in productos_por_linea.items():
            # Encabezado de línea con texto blanco y fondo verde institucional
            linea_style = ParagraphStyle(
                'LineaTitle',
                parent=self.styles['Heading2'],
                fontSize=11,
                textColor=colors.white,
                backColor=colors.HexColor('#4F9A54'),
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                leftIndent=6,
                rightIndent=6,
                spaceAfter=6,
                spaceBefore=6
            )
            
            # Tabla de encabezado de línea (cell merged)
            header_data = [[Paragraph("LÍNEA ESTRATÉGICA", linea_style)]]
            header_table = Table(header_data, colWidths=[7*inch])
            header_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            self.story.append(header_table)
            
            # Descripción de la línea estratégica
            desc_linea_style = ParagraphStyle(
                'DescLinea',
                parent=self.styles['Normal'],
                fontSize=10,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                spaceAfter=6,
                spaceBefore=6
            )
            desc_table = Table([[Paragraph(linea.upper(), desc_linea_style)]], colWidths=[7*inch])
            desc_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            
            self.story.append(desc_table)
            self.story.append(Spacer(1, 0.1*inch))
            
            # Estilo para encabezados con texto blanco
            white_header = ParagraphStyle('WhiteHeader', parent=self.styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold', fontSize=9)
            
            # Tabla de productos
            data = [[
                Paragraph('<b>PRODUCTO(S)</b>', white_header),
                Paragraph('<b>INDICADOR DE PRODUCTO</b>', white_header),
                Paragraph('<b>AVANCE DEL PRODUCTO</b>', white_header),
                Paragraph('<b>AVANCE FINANCIERO</b>', white_header)
            ]]
            
            for prod in productos:
                producto_text = prod.producto_mga or prod.codigo_producto
                indicador_text = prod.indicador_producto_mga or prod.personalizacion_indicador or 'N/A'
                
                # Calcular avance físico usando nuestra función
                avance_fisico_porcentaje = self.calcular_avance_producto(prod)
                avance_fisico = f"{avance_fisico_porcentaje:.1f}%"
                
                # Calcular avance financiero REAL desde ejecución presupuestal
                avance_financiero_porcentaje = self.calcular_avance_financiero(prod)
                avance_financiero = f"{avance_financiero_porcentaje:.1f}%"
                
                # Crear estilo justificado para textos largos
                justify_cell = ParagraphStyle(
                    'JustifyCell',
                    parent=self.styles['Normal'],
                    alignment=TA_JUSTIFY,
                    fontSize=8,
                    leading=10
                )
                
                data.append([
                    Paragraph(producto_text, justify_cell),  # Texto justificado completo
                    Paragraph(indicador_text, justify_cell),  # Texto justificado completo
                    Paragraph(avance_fisico, self.styles['Normal']),
                    Paragraph(avance_financiero, self.styles['Normal'])
                ])
            
            table = Table(data, colWidths=[2.5*inch, 2.5*inch, 1*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            
            self.story.append(table)
            self.story.append(Spacer(1, 0.15*inch))
    
    def generate_tabla_productos_por_sector(self):
        """Genera tabla detallada de productos por sector MGA"""
        # Agrupar productos por sector
        productos_por_sector = {}
        for prod in self.productos:
            sector = prod.sector_mga or 'Sin Sector'
            if sector not in productos_por_sector:
                productos_por_sector[sector] = []
            productos_por_sector[sector].append(prod)
        
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Paragraph(
            "DESCRIPCIÓN DE CUMPLIMIENTO DE METAS POR SECTORES MGA",
            title_style
        ))
        self.story.append(Spacer(1, 0.2*inch))
        
        for sector, productos in productos_por_sector.items():
            # Encabezado de sector
            sector_style = ParagraphStyle(
                'SectorTitle',
                parent=self.styles['Heading2'],
                fontSize=11,
                textColor=colors.white,
                backColor=colors.HexColor('#4F9A54'),
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                leftIndent=6,
                rightIndent=6,
                spaceAfter=6,
                spaceBefore=6
            )
            
            header_data = [[Paragraph("SECTOR MGA", sector_style)]]
            header_table = Table(header_data, colWidths=[7*inch])
            header_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            self.story.append(header_table)
            
            # Descripción del sector
            desc_sector_style = ParagraphStyle(
                'DescSector',
                parent=self.styles['Normal'],
                fontSize=10,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                spaceAfter=6,
                spaceBefore=6
            )
            desc_table = Table([[Paragraph(sector.upper(), desc_sector_style)]], colWidths=[7*inch])
            desc_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            
            self.story.append(desc_table)
            self.story.append(Spacer(1, 0.1*inch))
            
            # Tabla de productos
            white_header = ParagraphStyle('WhiteHeader', parent=self.styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold', fontSize=9)
            
            data = [[
                Paragraph('<b>PRODUCTO(S)</b>', white_header),
                Paragraph('<b>INDICADOR DE PRODUCTO</b>', white_header),
                Paragraph('<b>AVANCE DEL PRODUCTO</b>', white_header),
                Paragraph('<b>AVANCE FINANCIERO</b>', white_header)
            ]]
            
            for prod in productos:
                producto_text = prod.producto_mga or prod.codigo_producto
                indicador_text = prod.indicador_producto_mga or prod.personalizacion_indicador or 'N/A'
                avance_fisico_porcentaje = self.calcular_avance_producto(prod)
                avance_financiero_porcentaje = self.calcular_avance_financiero(prod)
                
                data.append([
                    Paragraph(producto_text, self.get_justify_style()),
                    Paragraph(indicador_text, self.get_justify_style()),
                    Paragraph(f"{avance_fisico_porcentaje:.1f}%", self.styles['Normal']),
                    Paragraph(f"{avance_financiero_porcentaje:.1f}%", self.styles['Normal'])
                ])
            
            table = Table(data, colWidths=[2.5*inch, 2.5*inch, 1*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            
            self.story.append(table)
            self.story.append(Spacer(1, 0.15*inch))
    
    def generate_tabla_productos_por_ods(self):
        """Genera tabla detallada de productos por ODS"""
        # Agrupar productos por ODS
        productos_por_ods = {}
        for prod in self.productos:
            ods = prod.ods or 'Sin ODS Asignado'
            if ods not in productos_por_ods:
                productos_por_ods[ods] = []
            productos_por_ods[ods].append(prod)
        
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Paragraph(
            "DESCRIPCIÓN DE CUMPLIMIENTO DE METAS POR OBJETIVOS DE DESARROLLO SOSTENIBLE",
            title_style
        ))
        self.story.append(Spacer(1, 0.2*inch))
        
        for ods, productos in productos_por_ods.items():
            # Encabezado de ODS
            ods_style = ParagraphStyle(
                'ODSTitle',
                parent=self.styles['Heading2'],
                fontSize=11,
                textColor=colors.white,
                backColor=colors.HexColor('#4F9A54'),
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                leftIndent=6,
                rightIndent=6,
                spaceAfter=6,
                spaceBefore=6
            )
            
            header_data = [[Paragraph("OBJETIVO DE DESARROLLO SOSTENIBLE", ods_style)]]
            header_table = Table(header_data, colWidths=[7*inch])
            header_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            
            self.story.append(header_table)
            
            # Descripción del ODS
            desc_ods_style = ParagraphStyle(
                'DescODS',
                parent=self.styles['Normal'],
                fontSize=10,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
                spaceAfter=6,
                spaceBefore=6
            )
            desc_table = Table([[Paragraph(ods.upper(), desc_ods_style)]], colWidths=[7*inch])
            desc_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            
            self.story.append(desc_table)
            self.story.append(Spacer(1, 0.1*inch))
            
            # Tabla de productos
            white_header = ParagraphStyle('WhiteHeader', parent=self.styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold', fontSize=9)
            
            data = [[
                Paragraph('<b>PRODUCTO(S)</b>', white_header),
                Paragraph('<b>INDICADOR DE PRODUCTO</b>', white_header),
                Paragraph('<b>AVANCE DEL PRODUCTO</b>', white_header),
                Paragraph('<b>AVANCE FINANCIERO</b>', white_header)
            ]]
            
            for prod in productos:
                producto_text = prod.producto_mga or prod.codigo_producto
                indicador_text = prod.indicador_producto_mga or prod.personalizacion_indicador or 'N/A'
                avance_fisico_porcentaje = self.calcular_avance_producto(prod)
                avance_financiero_porcentaje = self.calcular_avance_financiero(prod)
                
                data.append([
                    Paragraph(producto_text, self.get_justify_style()),
                    Paragraph(indicador_text, self.get_justify_style()),
                    Paragraph(f"{avance_fisico_porcentaje:.1f}%", self.styles['Normal']),
                    Paragraph(f"{avance_financiero_porcentaje:.1f}%", self.styles['Normal'])
                ])
            
            table = Table(data, colWidths=[2.5*inch, 2.5*inch, 1*inch, 1*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ]))
            
            self.story.append(table)
            self.story.append(Spacer(1, 0.15*inch))
    
    def generate_seccion_sectores(self):
        """Genera sección de avance por sectores MGA"""
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Paragraph(
            "AVANCE DE CUMPLIMIENTO DE METAS DEL PLAN DE DESARROLLO POR SECTORES",
            title_style
        ))
        
        desc_text = """
        Los sectores constituyen las áreas temáticas de acción gubernamental mediante las cuales se organiza 
        la gestión pública municipal. Cada sector agrupa objetivos, metas y programas de inversión específicos 
        orientados a atender las necesidades y prioridades de la población en campos determinados del desarrollo 
        territorial. Esta clasificación sectorial permite una gestión integral y articulada de las políticas 
        públicas, facilitando el seguimiento y evaluación de resultados por áreas de intervención.
        <br/><br/>
        La organización sectorial corresponde a la Metodología General Ajustada (MGA) establecida por el 
        Departamento Nacional de Planeación (DNP) para la formulación y evaluación de proyectos de inversión 
        pública en Colombia.
        """
        
        justify_style = ParagraphStyle(
            'Justify',
            parent=self.styles['BodyText'],
            alignment=TA_JUSTIFY,
            fontSize=10,
            spaceAfter=12
        )
        
        self.story.append(Paragraph(desc_text, justify_style))
        self.story.append(Spacer(1, 0.1*inch))
        
        # Generar gráfica moderna de sectores
        self.generate_grafica_moderna_sectores()
        
        # Generar tablas de productos por sector
        self.generate_tabla_productos_por_sector()
        
        self.story.append(PageBreak())
    
    def generate_seccion_ods(self):
        """Genera sección de avance por Objetivos de Desarrollo Sostenible"""
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        self.story.append(Paragraph(
            "AVANCE DE CUMPLIMIENTO DE METAS DEL PLAN DE DESARROLLO POR OBJETIVOS DE DESARROLLO SOSTENIBLE (ODS)",
            title_style
        ))
        
        desc_text = f"""
        Los Objetivos de Desarrollo Sostenible (ODS) constituyen un conjunto de 17 objetivos globales establecidos 
        por la Asamblea General de las Naciones Unidas en 2015, como parte integral de la Agenda 2030 para el 
        Desarrollo Sostenible. Esta agenda representa un compromiso universal de los Estados miembros para erradicar 
        la pobreza, proteger el planeta y garantizar que todas las personas gocen de paz, prosperidad y bienestar.
        <br/><br/>
        La República de Colombia, en cumplimiento de sus compromisos internacionales, ha incorporado los ODS en sus 
        instrumentos de planeación nacional y territorial. El Plan de Desarrollo Municipal de {self.entity.name} se encuentra 
        alineado con estos objetivos globales, contribuyendo desde el ámbito local al cumplimiento de las metas 
        establecidas en la Agenda 2030.
        <br/><br/>
        La presente sección evidencia la articulación entre las líneas estratégicas y sectores del Plan de Desarrollo 
        Municipal con los Objetivos de Desarrollo Sostenible, demostrando el compromiso de la administración municipal 
        con el desarrollo sostenible del territorio y el bienestar de sus habitantes.
        """
        
        justify_style = ParagraphStyle(
            'Justify',
            parent=self.styles['BodyText'],
            alignment=TA_JUSTIFY,
            fontSize=10,
            spaceAfter=12
        )
        
        self.story.append(Paragraph(desc_text, justify_style))
        self.story.append(Spacer(1, 0.2*inch))
        
        # Generar gráfica moderna de ODS
        self.generate_grafica_moderna_ods()
        
        # Generar tablas de productos por ODS
        self.generate_tabla_productos_por_ods()
        
        self.story.append(PageBreak())
    
    def generate_tabla_productos_detallada(self):
        """Genera tabla institucional por producto (formato oficial Sora-Boyacá)"""
        from app.models.pdm import PdmActividadEvidencia
        
        title_style = ParagraphStyle(
            'SectionTitle',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        anio_vigencia = "EL CUATRIENIO 2024-2027" if self.anio == 0 else f"LA VIGENCIA {self.anio}"
        
        self.story.append(PageBreak())
        self.story.append(Paragraph(
            f"EJECUCIÓN DEL PLAN DE ACCIÓN - {anio_vigencia}",
            title_style
        ))
        self.story.append(Spacer(1, 0.3*inch))
        
        # Agrupar actividades por producto
        actividades_por_producto = defaultdict(list)
        for act in self.actividades:
            if self.anio == 0 or act.anio == self.anio:
                actividades_por_producto[act.codigo_producto].append(act)
        
        # Procesar cada producto (SIN LÍMITE - mejora implementada)
        white_style = ParagraphStyle('WhiteText', parent=self.styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold', fontSize=10)
        
        total_productos = len(self.productos)
        print(f"   📦 Procesando {total_productos} productos...")
        
        for idx, prod in enumerate(self.productos, 1):
            print(f"   📦 Procesando producto {idx}/{total_productos}: {prod.codigo_producto}")
            
            # 1. LÍNEA ESTRATÉGICA (encabezado verde)
            linea_header = [[Paragraph('LÍNEA ESTRATÉGICA', white_style)]]
            linea_nombre = prod.linea_estrategica or 'SIN LÍNEA'
            linea_header.append([Paragraph(f'<b>{linea_nombre.upper()}</b>', self.styles['Normal'])])
            
            linea_table = Table(linea_header, colWidths=[7*inch])
            linea_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
            ]))
            self.story.append(linea_table)
            self.story.append(Spacer(1, 0.02*inch))
            
            # 2. PRODUCTO(S), INDICADOR, AVANCE DEL PRODUCTO, AVANCE FINANCIERO
            producto_nombre = prod.producto_mga or prod.codigo_producto
            indicador_nombre = prod.indicador_producto_mga or prod.personalizacion_indicador or 'N/A'
            avance_fisico = self.calcular_avance_producto(prod)
            avance_financiero = self.calcular_avance_financiero(prod)
            
            producto_data = [[
                Paragraph('<b>PRODUCTO(S)</b>', white_style),
                Paragraph('<b>INDICADOR DE PRODUCTO</b>', white_style),
                Paragraph('<b>AVANCE DEL PRODUCTO</b>', white_style),
                Paragraph('<b>AVANCE FINANCIERO</b>', white_style)
            ], [
                Paragraph(producto_nombre, self.styles['Normal']),
                Paragraph(indicador_nombre, self.styles['Normal']),
                Paragraph(f'<b>{avance_fisico:.0f}%</b>', self.styles['Normal']),
                Paragraph(f'<b>{avance_financiero:.0f}%</b>', self.styles['Normal'])
            ]]
            
            producto_table = Table(producto_data, colWidths=[2*inch, 2.5*inch, 1.25*inch, 1.25*inch])
            producto_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#C6EBBE')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            self.story.append(producto_table)
            self.story.append(Spacer(1, 0.02*inch))
            
            # 3. EJECUCIÓN PLAN DE ACCIÓN VIGENCIA
            actividades = actividades_por_producto.get(prod.codigo_producto, [])
            anio_vigencia = "2025" if self.anio == 2025 else str(self.anio) if self.anio > 0 else "2024-2027"
            
            ejecucion_header = [[Paragraph(f'<b>EJECUCIÓN PLAN DE ACCIÓN VIGENCIA {anio_vigencia}</b>', white_style)]]
            ejecucion_table = Table(ejecucion_header, colWidths=[7*inch])
            ejecucion_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#C6EBBE')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            self.story.append(ejecucion_table)
            self.story.append(Spacer(1, 0.02*inch))
            
            # 4. META Y/O ACTIVIDADES | INFORME DE EJECUCIÓN
            if actividades:
                total_actividades = len(actividades)
                completadas = sum(1 for act in actividades if act.estado == 'COMPLETADA')
                
                # Mostrar hasta 10 actividades con nombres completos (sin cortar)
                meta_text = f"La Meta No. <b>{total_actividades}</b> cuenta con la(s) siguiente(s) Actividades:<br/>"
                actividades_mostrar = actividades[:10]
                for idx, act in enumerate(actividades_mostrar, 1):
                    # Mostrar nombre completo sin recortar
                    nombre_actividad = act.nombre
                    meta_text += f"<b>{idx}.</b> {nombre_actividad}<br/>"
                
                if len(actividades) > 10:
                    meta_text += f"<i>... y {len(actividades) - 10} actividades más</i><br/>"
                
                informe_text = f"Se da cumplimiento a la meta con la ejecución de la siguiente contratación:<br/>"
                informe_text += f"<b>Total actividades:</b> {total_actividades}<br/>"
                informe_text += f"<b>Completadas:</b> {completadas}<br/>"
                if actividades[0].descripcion:
                    # Mostrar descripción completa sin recortar
                    informe_text += f"{actividades[0].descripcion}"
            else:
                meta_text = "Sin actividades registradas"
                informe_text = "No hay información de ejecución disponible"
            
            # Crear estilos justificados para contenido
            justify_style = ParagraphStyle(
                'JustifyContent',
                parent=self.styles['Normal'],
                alignment=TA_JUSTIFY,
                fontSize=9,
                leading=12
            )
            
            actividades_data = [[
                Paragraph('<b>Meta y/o Actividades</b>', white_style),
                Paragraph('<b>Informe de Ejecución</b>', white_style)
            ], [
                Paragraph(meta_text, justify_style),  # Texto justificado
                Paragraph(informe_text, justify_style)  # Texto justificado
            ]]
            
            # splitByRow=True permite que la tabla se divida entre páginas
            actividades_table = Table(actividades_data, colWidths=[3.5*inch, 3.5*inch], splitByRow=True)
            actividades_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Títulos centrados
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),  # Contenido justificado (via Paragraph)
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ]))
            self.story.append(actividades_table)
            self.story.append(Spacer(1, 0.02*inch))
            
            # 5. CANTIDAD META FÍSICA | RECURSOS | RESPONSABLE
            if self.anio == 0:
                # Cuatrienio completo: suma de todos los años
                total_recursos = (prod.total_2024 or 0) + (prod.total_2025 or 0) + (prod.total_2026 or 0) + (prod.total_2027 or 0)
            else:
                # Año específico
                if self.anio == 2024:
                    total_recursos = prod.total_2024 or 0
                elif self.anio == 2025:
                    total_recursos = prod.total_2025 or 0
                elif self.anio == 2026:
                    total_recursos = prod.total_2026 or 0
                elif self.anio == 2027:
                    total_recursos = prod.total_2027 or 0
                else:
                    total_recursos = 0
            
            responsable = prod.responsable_secretaria.nombre if prod.responsable_secretaria else 'N/A'
            
            recursos_data = [[
                Paragraph('<b>CANTIDAD DE META<br/>FÍSICA PROGRAMADA</b>', white_style),
                Paragraph('<b>Recursos Ejecutados</b>', white_style),
                Paragraph('<b>RESPONSABLE</b>', white_style)
            ], [
                Paragraph(f'<b>{prod.meta_cuatrienio or 1}</b>', self.styles['Normal']),
                Paragraph(f'<b>${total_recursos:,.0f}</b>', self.styles['Normal']),
                Paragraph(f'<b>{responsable.upper()}</b>', ParagraphStyle('Small', parent=self.styles['Normal'], fontSize=7))
            ]]
            
            recursos_table = Table(recursos_data, colWidths=[2.33*inch, 2.33*inch, 2.34*inch])
            recursos_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F9A54')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            self.story.append(recursos_table)
            self.story.append(Spacer(1, 0.02*inch))
            
            # 6. REGISTRO DE EVIDENCIA + IMÁGENES (OPTIMIZADO - carga bajo demanda)
            evidencias_encontradas = False
            
            # OPTIMIZACIÓN: Solo verificar si hay evidencias (sin cargar imágenes aún)
            actividades_con_evidencia = [act for act in actividades if hasattr(act, 'tiene_evidencia') and act.tiene_evidencia]
            
            if actividades_con_evidencia and self.db:
                # Cargar imágenes SOLO de las actividades con evidencia (query selectiva)
                from app.models.pdm import PdmActividadEvidencia
                actividades_ids_con_evidencia = [act.id for act in actividades_con_evidencia]
                
                # Query selectiva: solo imagenes de actividades con evidencia
                evidencias_completas = self.db.query(PdmActividadEvidencia).filter(
                    PdmActividadEvidencia.actividad_id.in_(actividades_ids_con_evidencia),
                    PdmActividadEvidencia.imagenes.isnot(None)
                ).all()
                
                # Crear diccionario de evidencias por actividad_id
                evidencias_dict = {ev.actividad_id: ev for ev in evidencias_completas}
                
                # Filtrar actividades que tienen imágenes
                evidencias_con_imagenes = [
                    act for act in actividades_con_evidencia 
                    if act.id in evidencias_dict and evidencias_dict[act.id].imagenes
                ]
                
                if evidencias_con_imagenes:
                    evidencias_encontradas = True
                    evidencia_header = [[Paragraph('<b>REGISTRO DE EVIDENCIAS</b>', white_style)]]
                    evidencia_table = Table(evidencia_header, colWidths=[7*inch])
                    evidencia_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F9A54')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Título centrado
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ('LEFTPADDING', (0, 0), (-1, -1), 10),
                    ]))
                    self.story.append(evidencia_table)
                    self.story.append(Spacer(1, 0.05*inch))
                    
                    # Procesar TODAS las evidencias
                    for num_evidencia, actividad in enumerate(evidencias_con_imagenes, 1):
                        evidencia = evidencias_dict[actividad.id]  # Obtener evidencia del dict
                        
                        # Subtítulo por actividad
                        actividad_nombre = actividad.nombre[:80] if len(actividad.nombre) > 80 else actividad.nombre
                        self.story.append(Paragraph(
                            f"<b>Actividad {num_evidencia}:</b> {actividad_nombre}",
                            ParagraphStyle('EvidenciaTitle', parent=self.styles['Normal'], fontSize=9, textColor=colors.HexColor('#003366'))
                        ))
                        self.story.append(Spacer(1, 0.05*inch))
                        
                        # Imágenes en grid 2x2 (sin límite de 4, pero paginadas)
                        if evidencia.imagenes and isinstance(evidencia.imagenes, list):
                            imagenes_cargadas = []
                            for idx, img_base64 in enumerate(evidencia.imagenes):  # SIN LÍMITE
                                try:
                                    if img_base64.startswith('data:image'):
                                        img_base64 = img_base64.split(',')[1]
                                    
                                    img_data = base64.b64decode(img_base64)
                                    
                                    # Tamaño optimizado: 3.3x3.3 pulgadas para grid 2x2
                                    img = RLImage(BytesIO(img_data), width=3.3*inch, height=3.3*inch, kind='proportional')
                                    imagenes_cargadas.append(img)
                                    print(f"      ✅ Evidencia {num_evidencia} - Imagen {idx+1} agregada")
                                except Exception as e:
                                    print(f"      ⚠️ Error evidencia {num_evidencia} imagen {idx+1}: {e}")
                        
                        
                        # Organizar imágenes en grid 2x2
                        if imagenes_cargadas:
                            grid_data = []
                            for i in range(0, len(imagenes_cargadas), 2):
                                row = imagenes_cargadas[i:i+2]
                                # Rellenar con celda vacía si es impar
                                if len(row) == 1:
                                    row.append('')
                                grid_data.append(row)
                            
                            img_table = Table(grid_data, colWidths=[3.5*inch, 3.5*inch])
                            img_table.setStyle(TableStyle([
                                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                ('TOPPADDING', (0, 0), (-1, -1), 5),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                            ]))
                            
                            self.story.append(img_table)
                            self.story.append(Spacer(1, 0.1*inch))
            
            if not evidencias_encontradas:
                evidencia_header = [[Paragraph('<b>REGISTRO DE EVIDENCIA</b>', white_style)]]
                evidencia_table = Table(evidencia_header, colWidths=[7*inch])
                evidencia_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F9A54')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Título centrado
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('TOPPADDING', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ]))
                self.story.append(evidencia_table)
            
            # Separador entre productos
            self.story.append(Spacer(1, 0.1*inch))
            self.story.append(PageBreak())
    
    def generate(self) -> bytes:
        """Genera el PDF completo y retorna los bytes"""
        try:
            print("📄 Generando informe PDM en PDF...")
            
            # Configurar documento
            self.doc = SimpleDocTemplate(
                self.buffer,
                pagesize=letter,
                rightMargin=0.5*inch,
                leftMargin=0.5*inch,
                topMargin=0.8*inch,
                bottomMargin=0.8*inch
            )
            
            # Estilos
            self.styles = getSampleStyleSheet()
            
            # 1. Portada
            print("  ├─ Portada")
            self.generate_portada()
            
            # 2. Introducción
            print("  ├─ Introducción")
            self.generate_introduccion()
            
            # 3. Sección de líneas estratégicas
            print("  ├─ Líneas Estratégicas")
            self.generate_seccion_lineas()
            
            # 4. Sección de sectores MGA
            print("  ├─ Sectores MGA")
            self.generate_seccion_sectores()
            
            # 5. Sección de ODS
            print("  ├─ Objetivos de Desarrollo Sostenible")
            self.generate_seccion_ods()
            
            # 6. Tabla de productos (versión básica - resumen)
            print("  ├─ Tabla de Productos (Resumen)")
            self.generate_tabla_productos()
            
            # 7. Tabla detallada de productos con actividades y evidencias
            print("  ├─ Detalle de Productos con Actividades y Evidencias")
            self.generate_tabla_productos_detallada()
            
            # Build PDF
            print("  └─ Construyendo PDF...")
            self.doc.build(
                self.story,
                onFirstPage=self.add_header_footer,
                onLaterPages=self.add_header_footer
            )
            
            pdf_bytes = self.buffer.getvalue()
            self.buffer.close()
            
            print(f"✅ PDF generado exitosamente ({len(pdf_bytes)} bytes)")
            return pdf_bytes
            
        except Exception as e:
            print(f"❌ Error generando PDF: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def generate_docx(self) -> bytes:
        """
        Genera informe en formato DOCX (Word)
        Nota: Requiere instalación de python-docx
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            print("📝 Generando informe DOCX...")
            
            doc = Document()
            
            # PORTADA
            anio_texto = "Todos los Años (2024-2027)" if self.anio == 0 else str(self.anio)
            doc.add_heading(f'INFORME DE GESTIÓN {anio_texto}', 0)
            doc.add_heading('PLAN DE DESARROLLO MUNICIPAL', 1)
            doc.add_heading(self.entity.name, 2)
            doc.add_page_break()
            
            # RESUMEN POR LÍNEAS ESTRATÉGICAS
            doc.add_heading('AVANCE POR LÍNEAS ESTRATÉGICAS', 1)
            
            # Agrupar por línea
            lineas_data = {}
            for prod in self.productos:
                linea = prod.linea_estrategica or 'Sin Línea'
                if linea not in lineas_data:
                    lineas_data[linea] = {'total': 0, 'suma_avance': 0}
                lineas_data[linea]['total'] += 1
                lineas_data[linea]['suma_avance'] += self.calcular_avance_producto(prod)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Línea Estratégica'
            hdr_cells[1].text = 'Avance (%)'
            
            for linea, data in lineas_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = linea
                promedio = data['suma_avance'] / data['total'] if data['total'] > 0 else 0
                row_cells[1].text = f"{promedio:.1f}%"
            
            doc.add_page_break()
            
            # TABLA DE PRODUCTOS CON MÁS DETALLES (mejora implementada)
            doc.add_heading('PRODUCTOS Y AVANCES DETALLADOS', 1)
            
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Código'
            hdr_cells[1].text = 'Producto'
            hdr_cells[2].text = 'Indicador'
            hdr_cells[3].text = 'Meta'
            hdr_cells[4].text = 'Avance Físico'
            hdr_cells[5].text = 'Avance Financiero'
            hdr_cells[6].text = 'Responsable'
            
            for prod in self.productos:
                row_cells = table.add_row().cells
                row_cells[0].text = prod.codigo_producto
                row_cells[1].text = (prod.producto_mga or prod.codigo_producto)[:100]
                row_cells[2].text = (prod.indicador_producto_mga or 'N/A')[:100]
                row_cells[3].text = str(prod.meta_cuatrienio or 0)
                row_cells[4].text = f"{self.calcular_avance_producto(prod):.1f}%"
                row_cells[5].text = f"{self.calcular_avance_financiero(prod):.1f}%"
                row_cells[6].text = prod.responsable_secretaria.nombre if prod.responsable_secretaria else 'N/A'
            
            doc.add_page_break()
            
            # SECCIÓN DE ACTIVIDADES POR PRODUCTO (nueva)
            doc.add_heading('ACTIVIDADES POR PRODUCTO', 1)
            
            actividades_por_producto = {}
            for act in self.actividades:
                if self.anio == 0 or act.anio == self.anio:
                    if act.codigo_producto not in actividades_por_producto:
                        actividades_por_producto[act.codigo_producto] = []
                    actividades_por_producto[act.codigo_producto].append(act)
            
            for prod in self.productos[:20]:  # Primeros 20 para no sobrecargar
                actividades = actividades_por_producto.get(prod.codigo_producto, [])
                if actividades:
                    doc.add_heading(f'{prod.codigo_producto} - {(prod.producto_mga or "")[:80]}', 2)
                    
                    act_table = doc.add_table(rows=1, cols=4)
                    act_table.style = 'Light List Accent 1'
                    act_hdr = act_table.rows[0].cells
                    act_hdr[0].text = 'Actividad'
                    act_hdr[1].text = 'Estado'
                    act_hdr[2].text = 'Meta'
                    act_hdr[3].text = 'Evidencia'
                    
                    for act in actividades[:10]:  # Max 10 por producto
                        act_row = act_table.add_row().cells
                        act_row[0].text = act.nombre[:100]
                        act_row[1].text = act.estado
                        act_row[2].text = str(act.meta_ejecutar or 0)
                        act_row[3].text = '✓ Sí' if (hasattr(act, 'tiene_evidencia') and act.tiene_evidencia) else '✗ No'
            
            # Guardar en BytesIO
            from io import BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_bytes = docx_buffer.getvalue()
            docx_buffer.close()
            
            print(f"✅ DOCX generado exitosamente ({len(docx_bytes)} bytes)")
            return docx_bytes
            
        except ImportError:
            print("❌ ERROR: python-docx no instalado")
            raise Exception("El formato DOCX no está disponible. Instale python-docx")
        except Exception as e:
            print(f"❌ Error generando DOCX: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def generate_excel(self) -> bytes:
        """
        Genera informe en formato Excel (XLSX)
        Usa openpyxl para crear un archivo Excel estructurado
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill
            from openpyxl.utils import get_column_letter
            
            print("📊 Generando informe Excel...")
            
            wb = Workbook()
            
            # HOJA 1: Resumen General
            ws = wb.active
            ws.title = "Resumen General"
            
            # Título
            ws['A1'] = f"INFORME PDM - {self.entity.name}"
            ws['A1'].font = Font(size=16, bold=True)
            anio_texto = "Todos los Años (2024-2027)" if self.anio == 0 else str(self.anio)
            ws['A2'] = f"Año: {anio_texto}"
            ws['A2'].font = Font(size=12)
            
            # Líneas Estratégicas
            ws['A4'] = "AVANCE POR LÍNEAS ESTRATÉGICAS"
            ws['A4'].font = Font(size=14, bold=True)
            
            ws['A5'] = "Línea Estratégica"
            ws['B5'] = "Avance (%)"
            ws['A5'].font = Font(bold=True)
            ws['B5'].font = Font(bold=True)
            
            row = 6
            lineas_data = {}
            for prod in self.productos:
                linea = prod.linea_estrategica or 'Sin Línea'
                if linea not in lineas_data:
                    lineas_data[linea] = {'total': 0, 'suma_avance': 0}
                lineas_data[linea]['total'] += 1
                lineas_data[linea]['suma_avance'] += self.calcular_avance_producto(prod)
            
            for linea, data in lineas_data.items():
                ws[f'A{row}'] = linea
                promedio = data['suma_avance'] / data['total'] if data['total'] > 0 else 0
                ws[f'B{row}'] = f"{promedio:.1f}%"
                row += 1
            
            # HOJA 2: Productos Detallados
            ws2 = wb.create_sheet("Productos")
            ws2['A1'] = "PRODUCTOS Y AVANCES"
            ws2['A1'].font = Font(size=14, bold=True)
            
            headers = ['Código', 'Producto', 'Indicador', 'Meta', 'Unidad', 'Avance Físico', 'Avance Financiero', 'Responsable']
            for col, header in enumerate(headers, 1):
                cell = ws2.cell(row=3, column=col)
                cell.value = header
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="003366", end_color="003366", fill_type="solid")
                cell.font = Font(bold=True, color="FFFFFF")
            
            row = 4
            for prod in self.productos:
                ws2[f'A{row}'] = prod.codigo_producto
                ws2[f'B{row}'] = prod.producto_mga or 'N/A'
                ws2[f'C{row}'] = prod.indicador_producto_mga or 'N/A'
                ws2[f'D{row}'] = prod.meta_cuatrienio or 0
                ws2[f'E{row}'] = prod.unidad_medida or ''
                ws2[f'F{row}'] = f"{self.calcular_avance_producto(prod):.1f}%"
                ws2[f'G{row}'] = f"{self.calcular_avance_financiero(prod):.1f}%"
                ws2[f'H{row}'] = prod.responsable_secretaria.nombre if prod.responsable_secretaria else 'N/A'
                row += 1
            
            # Ajustar anchos de columna
            for col in range(1, 9):
                ws2.column_dimensions[get_column_letter(col)].width = 20
            
            # HOJA 3: Actividades (mejora implementada)
            ws3 = wb.create_sheet("Actividades")
            ws3['A1'] = "ACTIVIDADES Y ESTADOS"
            ws3['A1'].font = Font(size=14, bold=True)
            
            headers_act = ['Código Producto', 'Actividad', 'Estado', 'Año', 'Meta Ejecutar', 'Fecha Inicio', 'Fecha Fin', 'Responsable', 'Evidencia']
            for col, header in enumerate(headers_act, 1):
                cell = ws3.cell(row=3, column=col)
                cell.value = header
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="4F9A54", end_color="4F9A54", fill_type="solid")
                cell.font = Font(bold=True, color="FFFFFF")
            
            row = 4
            for act in self.actividades:
                if self.anio == 0 or act.anio == self.anio:
                    ws3[f'A{row}'] = act.codigo_producto
                    ws3[f'B{row}'] = act.nombre[:100]
                    ws3[f'C{row}'] = act.estado
                    ws3[f'D{row}'] = act.anio
                    ws3[f'E{row}'] = act.meta_ejecutar or 0
                    ws3[f'F{row}'] = act.fecha_inicio.strftime('%Y-%m-%d') if act.fecha_inicio else ''
                    ws3[f'G{row}'] = act.fecha_fin.strftime('%Y-%m-%d') if act.fecha_fin else ''
                    ws3[f'H{row}'] = act.responsable_secretaria.nombre if act.responsable_secretaria else 'N/A'
                    ws3[f'I{row}'] = 'Sí' if (hasattr(act, 'tiene_evidencia') and act.tiene_evidencia) else 'No'
                    row += 1
            
            # Ajustar anchos
            for col in range(1, 10):
                ws3.column_dimensions[get_column_letter(col)].width = 18
            
            # HOJA 4: Estadísticas de Evidencias (mejora implementada)
            ws4 = wb.create_sheet("Evidencias")
            ws4['A1'] = "RESUMEN DE EVIDENCIAS"
            ws4['A1'].font = Font(size=14, bold=True)
            
            ws4['A3'] = 'Métrica'
            ws4['B3'] = 'Valor'
            ws4['A3'].font = Font(bold=True)
            ws4['B3'].font = Font(bold=True)
            
            total_actividades = len([a for a in self.actividades if self.anio == 0 or a.anio == self.anio])
            actividades_con_evidencia = len([a for a in self.actividades if (self.anio == 0 or a.anio == self.anio) and a.evidencia])
            porcentaje_evidencia = (actividades_con_evidencia / total_actividades * 100) if total_actividades > 0 else 0
            
            ws4['A4'] = 'Total Actividades'
            ws4['B4'] = total_actividades
            ws4['A5'] = 'Actividades con Evidencia'
            ws4['B5'] = actividades_con_evidencia
            ws4['A6'] = 'Porcentaje Documentado'
            ws4['B6'] = f"{porcentaje_evidencia:.1f}%"
            
            ws4['A8'] = 'Productos'
            ws4['B8'] = len(self.productos)
            ws4['A9'] = 'Avance Físico Promedio'
            suma_avances = sum(self.calcular_avance_producto(p) for p in self.productos)
            ws4['B9'] = f"{suma_avances / len(self.productos):.1f}%" if self.productos else "0%"
            
            # Guardar en BytesIO
            from io import BytesIO
            excel_buffer = BytesIO()
            wb.save(excel_buffer)
            excel_bytes = excel_buffer.getvalue()
            excel_buffer.close()
            
            print(f"✅ Excel generado exitosamente con 4 hojas ({len(excel_bytes)} bytes)")
            return excel_bytes
            
        except ImportError as ie:
            print(f"❌ ERROR: Biblioteca no instalada - {ie}")
            raise Exception("El formato Excel no está disponible. Instale openpyxl")
        except Exception as e:
            print(f"❌ Error generando Excel: {e}")
            import traceback
            traceback.print_exc()
            raise
